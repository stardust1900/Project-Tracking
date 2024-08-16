from datetime import datetime
import re
from docx import Document
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from models import Dept, Log, Product, Project, User
def get_between_chars(s, start_char, end_char):
    pattern = f'{re.escape(start_char)}(.+?){re.escape(end_char)}'
    # print(pattern)
    result = re.search(pattern, s)
    if result:
        return result.group(1)
    else:
        return None
    
def read_docx(doc,product_m,dept_m,user_m,session):
    products=[]
    managers = []
    supporters = []
    depts = []
    # 遍历文档中的段落
    product = None
    project = None
    instruction = None
    dept = None
    manager = None
    input_date = None
    supporter_list = []
    logs = []

    for para in doc.paragraphs:
        line = para.text.strip()
        if line.startswith('【'):
            if project:
                print('product:%s project:%s dept:%s manager:%s input_date:%s' % (product,project,dept,manager,input_date))
                if not instruction:
                    instruction = ''
                print(instruction)
                print(logs)
                if not supporter_list:
                    supporter_list.append('无')
                save(product=product,project=project,dept=dept,manager=manager,input_date=input_date,instruction=instruction,supporter_list=supporter_list,
                     logs=logs,product_m=product_m,dept_m=dept_m,user_m=user_m)
                
                product = None
                project = None
                instruction = None
                dept = None
                manager = None
                input_date = None
                supporter_list = []
                logs = []
                # break
            product = get_between_chars(line,'【','】')
            if product not in products:
                products.append(product)
            # print(product)
            # project = get_between_chars(line,'】','(')
            # if not project:
            project = get_between_chars(line,'】','（')
            # print(project)
            # dept_m_d = get_between_chars(line,'(',')')
            # if not dept_m_d:
            dept_m_d = get_between_chars(line,'（','）')
            # print(dept_m_d)
            array = dept_m_d.split(' ')
            # print(array)
            dept = array[0]
            manager = array[1].strip()
            if not manager:
                print(dept_m_d)
                break
            input_date = array[2]
            if dept not in depts:
                depts.append(dept)
            if manager not in managers:
                managers.append(manager)
            # print(dept_m_d)

            pattern = r'@(.*)'
            result = re.search(pattern, line)
            if result:
                supporter = result.group(1)
                if '@' in supporter:
                    spArray = supporter.split('@')
                    supporter_list = spArray
                else:
                    supporter_list.append(supporter)
                for s in supporter_list:
                    if s not in supporters:
                        supporters.append(s)
            # print(supporter)
        elif line.startswith('2024'): #日志
            if project == '蓝海银行':
                print('log: %s' % line)
            # print(line)
            logs.append(line)
        else: # 项目说明
            if product and not logs:
                instruction = line
                # print(instruction)
    
    if project:
                print('product:%s project:%s dept:%s manager:%s input_date:%s' % (product,project,dept,manager,input_date))
                print(instruction)
                print(logs)
                save(product=product,project=project,dept=dept,manager=manager,input_date=input_date,instruction=instruction,supporter_list=supporter_list,
                     logs=logs,product_m=product_m,dept_m=dept_m,user_m=user_m)
    # print('product:%s project:%s dept:%s manager:%s input_date:%s' % (product,project,dept,manager,input_date))
    # print(instruction)
    # print(logs)
    print(products)
    # print(managers)
    # print(supporters)
    # print(depts)
    # for dept in depts:
    #     print("insert into dept(name,create_time) values ('%s',now());" % dept)



def save(product,project,dept,manager,input_date,instruction,supporter_list,logs,product_m,dept_m,user_m):

    p = Project(merchant_name=project,product_id=product_m[product].id,instruction=instruction,
                input_date=input_date,dept_id=dept_m[dept].id,manager_id=user_m[manager].id,create_time=datetime.today())
    for s in supporter_list:
        p.supporters.append(user_m[s])
    print(p)
    session.add(p)
    session.commit()
    print(p)
    for log in logs:
        log_date = log[0:10]
        content = log[10:]
        l = Log(project_id=p.id,content=content,user_id=user_m[supporter_list[0]].id,input_date=log_date,create_time=datetime.today())
        print(l)
        session.add(l)
    session.commit()

if __name__ == "__main__":
    doc = Document('project.docx')
    engine = create_engine('mysql://projectuser:123456@localhost/project')
    Session = sessionmaker(bind=engine)
    session = Session()
    product_m = {}
    products = session.query(Product).all()
    for p in products:
        product_m[p.name] =p
    # print(product_m)
    depts = session.query(Dept).all()
    dept_m = {}
    for d in depts:
        dept_m[d.name] = d
    # print(dept_m)
    user_m = {}
    users = session.query(User).all()
    for u in users:
        user_m[u.name] = u
    # print(user_m)

    read_docx(doc,product_m=product_m,dept_m=dept_m,user_m=user_m,session=session)
# i = 0
# headings = []
# for para in doc.paragraphs:
#     i = i+1
#     print(para.style)
#     print(para.text)
#     if para.style.name.startswith('Heading'):
#         headings.append((para.style.name, para.text))
#     for run in para.runs:
#         # 获取字体名称
#         font_name = run.font.name
#         # 获取字体大小
#         font_size = run.font.size
#         # 获取字体颜色
#         font_color = run.font.color.rgb if run.font.color else None
#         print(f"Font Name: {font_name}, Font Size: {font_size}, Font Color: {font_color}")
#         print(run._element.xml)
#     if i >=10:
#         break

# print(headings)
