from datetime import datetime, timedelta
import functools
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side
from flask import Flask, flash, g, get_flashed_messages, redirect,render_template,request, send_file, url_for
import logging
from logging.handlers import TimedRotatingFileHandler
from flask_login import LoginManager, current_user, login_required, login_user, logout_user
from flask_migrate import Migrate
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import or_
from werkzeug.security import generate_password_hash, check_password_hash
from models import Dept, Log, Product, Project, db,User
import config

# 日志的等级的设置
logging.basicConfig(level=logging.DEBUG)
# 创建日志记录器，　指明日志保存的路径
file_log_handler = TimedRotatingFileHandler('./logs/pt.log','midnight',1,30,'utf-8')
# 设置日志的格式       日志等级       日志信息的文件名　　行数　　日志信息
formatter = logging.Formatter('%(levelname)s %(filename)s %(lineno)d %(message)s')
# 将日志记录器指定日志的格式
file_log_handler.setFormatter(formatter)
# 为全局的日志工具对象添加日志记录器
# logging.getLogger('sqlalchemy.engine').setLevel(logging.INFO)
logging.getLogger().addHandler(file_log_handler)


app = Flask(__name__)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
app.config['SECRET_KEY'] = config.SECRET_KEY
app.config['SQLALCHEMY_DATABASE_URI'] = config.DATABASE_URI
db.init_app(app)
migrate = Migrate(app, db)
with app.app_context():
    if config.DROP_TABLES:
        logging.info('开始删除旧表')
        db.drop_all()
        logging.info('旧表删除完成')
    db.create_all()

@app.route("/")
@login_required
def main():
    products = Product.query.filter(Product.level==0).all()
    projects = Project.query.all()
    depts = []
    managers = []
    supporters = []
    # print(projects)
    for p in projects:
        # print(p.dept,p.manager,p.supporters,p.product)
        if p.dept not in depts:
            depts.append(p.dept)
        if p.manager not in managers:
            managers.append(p.manager)
        for supporter in p.supporters:
            if supporter not in supporters:
                supporters.append(supporter)
        
    return render_template('main.html',products=products,projects=projects,depts=depts,managers=managers,supporters=supporters)

@app.route("/query")
def query():
    products = Product.query.filter(Product.level==0).all()
    projects = Project.query.all()
    depts = []
    managers = []
    supporters = []
    # print(projects)
    for p in projects:
        # print(p.dept,p.manager,p.supporters,p.product)
        if p.dept not in depts:
            depts.append(p.dept)
        if p.manager not in managers:
            managers.append(p.manager)
        for supporter in p.supporters:
            if supporter not in supporters:
                supporters.append(supporter)
    # logging.info(request.args)
    qProjectName = request.args.get('qProjectName',None)
    qDept = request.args.get('qDept',None)
    qManager = request.args.get('qManager',None)
    qSupporter = request.args.get('qSupporter',None)
    
    if qProjectName:
        projects = list(filter(lambda p:qProjectName in p.merchant_name,projects))
    if qDept:
        # logging.info("qDept:%s",qDept)
        qDept = int(qDept)
        # logging.info("qDept:%s",qDept)
        projects = list(filter(lambda p:p.dept_id == qDept,projects))
        # print("projects:",projects)
    if qManager:
        qManager = int(qManager)
        projects = list(filter(lambda p:p.manager_id == qManager,projects))
    if qSupporter:
        qSupporter = int(qSupporter)
        projects = list(filter(lambda p: any(s.id == qSupporter for s in p.supporters) ,projects))

    return render_template('main.html',products=products,projects=projects,depts=depts,managers=managers,supporters=supporters,qProjectName=qProjectName,qDept=qDept,qManager=qManager,qSupporter=qSupporter)

@app.route("/login", methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        login_name=request.form['login_name']
        password = request.form['password']
        # logging.info("login_name:%s password:%s",request.form['login_name'],request.form['password'])
        hashed_password = generate_password_hash(request.form['password'])
        # logging.info(hashed_password)
        user = User.query.filter(or_(User.email==login_name,User.login_name==login_name)).first()
        logging.info(user)
        if user:
            if check_password_hash(user.password,password):
                logging.info("验证成功")
                # logging.info(current_user)
                login_user(user)
                # logging.info(current_user)
                # return redirect(url_for('main'))
                return redirect(url_for('main'))
            else:
                flash('用户名或密码错误', 'warning')
                return render_template('login.html',)
        else:
            flash('用户不存在', 'warning')
            return render_template('login.html',)
    else:
        # 清空flash信息
        flash("")
        logging.info("flash:%s",get_flashed_messages())
        return render_template('login.html')

@login_manager.user_loader
def load_user(user_id):
    return User.query.filter_by(id=user_id).first()

@login_manager.unauthorized_handler
def unauthorized():
    flash('请先登录以访问此页面。', 'info')  # 修改登录消息
    return redirect(url_for('login'))

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/addProject', methods=['GET', 'POST'])
@login_required
def addProject():
    if request.method == 'POST':
        projectId = request.form['projectId']
        project = Project()
        if projectId:
            project = Project.query.get(projectId)
        productId=request.form['productId']
        merchantName = request.form['merchantName']
        instruction = request.form['instruction']
        input_date = request.form['input_date']
        launch_date = request.form['launch_date']
        deptId = request.form['deptId']
        manger = request.form['manger']
        supporters = request.form.getlist('supporters')
        projectState = request.form['projectState']
        remark = request.form['remark']
        logging.info('projectId:%s productId:%s merchantName:%s manger:%s supporters:%s projectState:%s' % (projectId, productId,merchantName,manger,supporters,projectState))

        project.product_id = productId
        project.merchant_name = merchantName
        project.instruction = instruction
        project.input_date = input_date
        project.launch_date = launch_date
        project.dept_id = deptId
        project.manager_id = manger
        project.supporters = []
        for s in supporters:
            u = User.query.get(s)
            if u and u not in project.supporters:
                project.supporters.append(u)
        project.state = projectState
        project.remark = remark
        if not project.create_time:
            project.create_time = datetime.now()
            project.create_by = current_user.id
        project.update_time = datetime.now()
        project.update_by = current_user.id
        db.session.merge(project)
        db.session.commit()
        return redirect(url_for('main'))
    else:
        productId = request.args.get('productId')
        if productId:
            product = Product.query.get(productId)
            project = Project()
            project.input_date = datetime.now().strftime('%Y%m%d')
            project.launch_date = datetime.now().strftime('%Y-%m-%d')
            depts = Dept.query.all()
            managers = User.query.all()
            supporters = User.query.filter(User.dept_id == 11).all()
            if product:
                return render_template('project.html',project=project,product=product,depts=depts,managers=managers,supporters=supporters)
            else:
                return redirect(url_for('main'))
        else:
            return redirect(url_for('main'))
        
@app.route('/editProject')
@login_required
def editProject():
    projectId = request.args.get('projectId',None)
    project = Project.query.get(projectId)
    depts = Dept.query.all()
    managers = User.query.all()
    supporters = User.query.filter(User.dept_id == 11).all()
    return render_template('project.html',project=project,product=project.product,depts=depts,managers=managers,supporters=supporters)

@app.route('/addLog', methods=['GET', 'POST'])
@login_required
def addLog():
    if request.method == 'POST':
        projectId = request.form['projectId']
        logId = request.form['logId']
        input_date = request.form['input_date']
        content = request.form['content']
        log = Log()
        if logId:
            log = Log.query.get(logId)
        log.project_id = projectId
        log.input_date = input_date
        log.content = content
        log.user_id = current_user.id
        if not log.create_time:
            log.create_time = datetime.now()
            log.create_by = current_user.id
        log.update_by = current_user.id
        log.update_time = datetime.now()
        db.session.merge(log)
        db.session.commit()
        return redirect(url_for('main'))
    else:
        projectId = request.args.get('projectId')
        project = Project.query.get(projectId)
        log = Log()
        log.input_date = datetime.now().strftime('%Y-%m-%d')
        return render_template('log.html',project=project,log=log)
    
@app.route('/editLog')
@login_required
def editLog():
    logId = request.args.get('logId')
    log = Log.query.get(logId)
    return render_template('log.html',project=log.project,log=log)

@app.route('/downloadWord')
@login_required
def downloadWord():

    nearest_monday = get_nearest_monday()
    logging.info('nearest_monday:%s',nearest_monday)

    nearest_logs = Log.query.filter(Log.create_time >= nearest_monday).all()
    project_ids = [log.project_id for log in nearest_logs]
    target_projects = Project.query.filter(Project.id.in_(project_ids)).distinct().all()
    products = Product.query.filter(Product.level==0).all()

    # 创建一个新的Word文档
    doc = Document()

    # 获取文档的默认样式，并将其字体设置为等线
    default_style = doc.styles['Normal']
    default_style.font.name = '等线'

    # 添加一个居中的段落
    paragraph = doc.add_paragraph()
    # paragraph.alignment = 1  # 设置段落居中对齐（1 表示居中）

    # 在段落中添加一个 Run 对象，并设置其字体大小
    run = paragraph.add_run('项目列表清单')
    run.font.size = Pt(26)  # 设置字体大小为 26 磅 一号字
    run.bold = True # 加粗字体
    # 添加一级标题  标题18 小二  内容11
    # title = doc.add_heading('项目列表清单', level=1)

    for i,product in enumerate(products,start=1):
        write_title(doc,num_to_chinese(i)+'、'+product.name)
        for idx, project in enumerate(target_projects,start=1):
            supporter_text = ''
            for supporter in project.supporters:
                supporter_text=supporter_text+'@'+supporter.name

            p_list = '【%s】 %s （%s %s %s）%s' % (project.product.name,project.merchant_name,project.dept.name,project.manager.name,project.input_date,supporter_text)
            write_list(doc,p_list,idx)
            # 项目说明
            doc.add_paragraph(project.instruction)
            last_log = project.logs[-1]
            doc.add_paragraph(last_log.input_date+' '+last_log.content)


    # 保存文档到内存中的字节流
    doc_stream = io.BytesIO()
    doc.save(doc_stream)

    # 将字节流的位置重置为开始
    doc_stream.seek(0)

    filename = '技术支持部周报%s.docx' % datetime.today().strftime('%Y-%m-%d')
    # 使用send_file方法发送文件
    return send_file(doc_stream, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

def write_title(doc,text):
    # 添加一个一级标题，并设置其字体为小二
    title = doc.add_heading(text, level=1)
    run = title.runs[0]
    # run.font.name = 'SimSun'  # 设置字体为宋体（小二）
    run.font.size = Pt(18)  # 设置字体大小为 18 磅

def write_list(doc,text,idx):
    p = doc.add_paragraph()
    run = p.add_run(f'{idx}. {text}')
    run.font.size = Pt(14)
    p_format = p.paragraph_format
    p_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p_format.numbering_level = 0
    p_format.first_line_indent = Pt(0)
    p_format.left_indent = Pt(18)
    p_format.numbering_start = idx

def get_nearest_monday():
    today = datetime.today()
    days_since_monday = today.weekday()  # 0 表示周一，1 表示周二，依此类推
    days_to_nearest_monday = (days_since_monday - 0) % 7
    days_to_nearest_monday = 7 if days_to_nearest_monday == 0 else days_to_nearest_monday
    nearest_monday = today - timedelta(days=days_to_nearest_monday)
    return nearest_monday.date()

def num_to_chinese(num):
    num_dict = {
        1: "一",
        2: "二",
        3: "三",
        4: "四",
        5: "五",
        6: "六",
        7: "七",
        8: "八",
        9: "九"
    }
    return num_dict.get(num, "")

@app.route('/downloadExcel')
@login_required
def downloadExcel():
    projects = Project.query.all()
    nearest_monday = get_nearest_monday().strftime('%Y-%m-%d')
    # list(filter(lambda p:qProjectName in p.merchant_name,projects))
    complete_projects = list(filter(lambda p: p.state in [3] and p.launch_date and p.launch_date < nearest_monday ,projects))
    tracking_projects = list(filter(lambda p: p.state in [0,1,2,4,5] or (p.state in [3] and (not p.launch_date or p.launch_date >= nearest_monday)),projects))
    stopped_projects = list(filter(lambda p: p.state in [6,7] ,projects))

    # logging.info('complete_projects:%s',complete_projects)
    # logging.info('tracking_projects:%s',tracking_projects)
    # logging.info('stopped_projects:%s',stopped_projects)
      # 创建一个 Workbook 对象
    wb = Workbook()
    
    # 创建第一个工作表
    ws1 = wb.create_sheet(title="已完成项目")
    headers = ['序号', '商户名称', '产品','产品详情','提交日期','完成日期','归属部门','客户经理','技术支持','备注']
    ws1.append(headers)
    ws1.column_dimensions['B'].width = 40
    ws1.column_dimensions['D'].width = 100
    row = []
    for i,p in enumerate(complete_projects,start=1):
        row.append(i)
        row.append(p.merchant_name)
        row.append(p.product.name)
        row.append(p.instruction)
        row.append(p.input_date)
        row.append(p.launch_date)
        row.append(p.dept.name)
        row.append(p.manager.name)
        spt = ''
        for s in p.supporters:
            spt = spt + s.name + ' '
        row.append(spt)
        row.append(p.remark)
        ws1.append(row)
        row.clear()

    # 创建第二个工作表
    ws2 = wb.create_sheet(title="跟进中项目")
    headers = ['序号', '商户名称', '产品','产品详情','提交日期','归属部门','客户经理','技术支持','项目近期进展','备注']
    ws2.append(headers)
    ws2.column_dimensions['B'].width = 40
    ws2.column_dimensions['D'].width = 100
    ws2.column_dimensions['I'].width = 100
    row = []
    reds=[]
    oranges=[]
    greens=[]
    blues=[]
    for i,p in enumerate(tracking_projects,start=1):
        row.append(i)
        row.append(p.merchant_name)
        row.append(p.product.name)
        row.append(p.instruction)
        row.append(p.input_date)
        row.append(p.dept.name)
        row.append(p.manager.name)
        spt = ''
        for s in p.supporters:
            spt = spt + s.name + ' '
        row.append(spt)
        if p.logs :
            row.append(p.logs[-1].content)
        else:
            row.append('')
        row.append(p.remark)
        ws2.append(row)
        row.clear()
        # 项目状态 0-初始 1-对接中 2-对接完成待上线 3-已上线 4-本周无进展 5-两周无进展 6-长期无进展 7-终止
        if p.state == 4:
            oranges.append(i+1)
        elif p.state == 5:
            reds.append(i+1)
        elif p.state == 3:
            greens.append(i+1)
        elif p.state == 2:
            blues.append(i+1) 
    
    change_color(ws2,oranges,'FFA500')
    change_color(ws2,reds,'FF0000')
    change_color(ws2,greens,'00FF00')
    change_color(ws2,blues,'0000FF')
    
    # 创建第三个工作表
    ws3 = wb.create_sheet(title="无进展项目")
    headers = ['序号', '商户名称', '产品','产品详情','提交日期','归属部门','客户经理','技术支持','备注']
    ws3.append(headers)
    ws3.column_dimensions['B'].width = 40
    ws3.column_dimensions['D'].width = 100
    ws3.column_dimensions['I'].width = 100
    row = []
    for i,p in enumerate(tracking_projects,start=1):
        row.append(i)
        row.append(p.merchant_name)
        row.append(p.product.name)
        row.append(p.instruction)
        row.append(p.input_date)
        row.append(p.dept.name)
        row.append(p.manager.name)
        spt = ''
        for s in p.supporters:
            spt = spt + s.name + ' '
        row.append(spt)
        row.append(p.remark)
        ws3.append(row)
        row.clear()

     # 删除默认的工作表
    wb.remove(wb['Sheet'])

    change_sheet_style(ws1)
    change_sheet_style(ws2)
    change_sheet_style(ws3)
    # 将 Workbook 对象保存为 Excel 文件
    excel_file_path = 'exported_data.xlsx'
    wb.save(excel_file_path)
    filename = '技术开展项目清单%s.xlsx' % datetime.today().strftime('%Y%m%d')
    return send_file(excel_file_path, as_attachment=True, download_name=filename)

def change_sheet_style(worksheet):
    header_font = Font(bold=True)
    header_fill = PatternFill(start_color='D3D3D3', end_color='D3D3D3', fill_type='solid')
    for cell in worksheet[1]:
        cell.font = header_font
        cell.fill = header_fill

    # 设置表格边框
    border = Border(left=Side(style='thin'), right=Side(style='thin'),
                    top=Side(style='thin'), bottom=Side(style='thin'))

    for row in worksheet.iter_rows(min_row=1, max_row=worksheet.max_row,
                                min_col=1, max_col=worksheet.max_column):
        for cell in row:
            cell.border = border

def change_color(worksheet,rownums,color):
    logging.info('rownums:%s color:%s',rownums,color)
    for rownum in rownums:
        c_font = Font(color=color)
        for cell in worksheet[rownum]:
            cell.font = c_font